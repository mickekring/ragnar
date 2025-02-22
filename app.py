
# Python imports
import os
import streamlit as st
import hashlib

# External imports
from docx import Document
from pydub import AudioSegment

# Local imports
from functions.functions import convert_to_mono_and_compress
from functions.transcribe import transcribe_with_kb_whisper, transcribe_with_whisper
import config as c


### INITIAL VARIABLES

# Creates folder if they don't exist
os.makedirs("audio", exist_ok=True) # Where audio/video files are stored for transcription
os.makedirs("text", exist_ok=True) # Where transcribed document are beeing stored


# Check and set default values if not set in session_state
# of Streamlit
if "translation" not in st.session_state: # If audio has been translated
    st.session_state["translation"] = False
if "cpu_vs_gpu" not in st.session_state: # If user device has GPU support
    st.session_state["cpu_vs_gpu"] = False
if "spoken_language" not in st.session_state: # What language source audio is in
    st.session_state["spoken_language"] = "Svenska"
if "transcribe_model" not in st.session_state: # What model of Whisper to use
    st.session_state["transcribe_model"] = "KB Whisper Small"
if "file_name_converted" not in st.session_state: # Audio file name
    st.session_state["file_name_converted"] = None


# Checking if uploaded or recorded audio file has been transcribed
def compute_file_hash(uploaded_file):

    # Compute the MD5 hash of a file
    hasher = hashlib.md5()
    
    for chunk in iter(lambda: uploaded_file.read(4096), b""):
        hasher.update(chunk)
    uploaded_file.seek(0)  # Reset the file pointer to the beginning
    
    return hasher.hexdigest()


### MAIN APP ###########################

# Page configuration
st.set_page_config(
    page_title="Ragnar",
    page_icon=None,
    layout="centered",
    initial_sidebar_state="auto"
    )


def main():

    global translation
    global model_map_transcribe_model

    ### SIDEBAR

    # Sidebar image of Ragnar
    st.sidebar.image("images/ragge3.png", width = 220)

    ###### SIDEBAR SETTINGS

    st.sidebar.header("Inställningar")
    st.sidebar.markdown("")

    # Dropdown menu - choose Whisper model
    transcribe_model = st.sidebar.selectbox(
            "Välj modell för transkribering", 
            [
            "KB Whisper Large", 
             "KB Whisper Medium", 
             "KB Whisper Small",
             "KB Whisper Base",
             "KB Whisper Tiny",
             "OpenAI Whisper Turbo",
             "OpenAI Whisper Large",
             "OpenAI Whisper Medium",
             "OpenAI Whisper Small",
             "OpenAI Whisper Base",
             "OpenAI Whisper Tiny"
             ],
            index=[
                "KB Whisper Large", 
                "KB Whisper Medium", 
                "KB Whisper Small",
                "KB Whisper Base",
                "KB Whisper Tiny",
                "OpenAI Whisper Turbo",
                "OpenAI Whisper Large",
                "OpenAI Whisper Medium",
                "OpenAI Whisper Small",
                "OpenAI Whisper Base",
                "OpenAI Whisper Tiny"
                ].index(st.session_state["transcribe_model"]),
        )

    model_map_transcribe_model = {
            "KB Whisper Large": "kb-whisper-large",
            "KB Whisper Medium": "kb-whisper-medium",
            "KB Whisper Small": "kb-whisper-small",
            "KB Whisper Base": "kb-whisper-base",
            "KB Whisper Tiny": "kb-whisper-tiny",
            "OpenAI Whisper Turbo": "turbo",
            "OpenAI Whisper Large": "large",
            "OpenAI Whisper Medium": "medium",
            "OpenAI Whisper Small": "small",
            "OpenAI Whisper Base": "base",
            "OpenAI Whisper Tiny": "tiny"
        }

    # Dropdown menu - choose source language of audio
    spoken_language = st.sidebar.selectbox(
            "Välj språk som talas", 
            ["Automatiskt", "Svenska", "Engelska", "Franska", "Tyska", "Spanska"],
            index=["Automatiskt", "Svenska", "Engelska", "Franska", "Tyska", "Spanska"].index(st.session_state["spoken_language"]),
        )

    model_map_spoken_language = {
            "Automatiskt": None,
            "Svenska": "sv",
            "Engelska": "en",
            "Franska": "fr",
            "Tyska": "de",
            "Spanska": "sp"

        }

    # Update the session_state directly
    st.session_state["transcribe_model"] = transcribe_model
    st.session_state["spoken_language"] = spoken_language

    print(model_map_transcribe_model[st.session_state["transcribe_model"]])
    print(model_map_spoken_language[st.session_state["spoken_language"]])

    st.sidebar.markdown(
        "#"
        )

    st.sidebar.markdown(f"""
    Version: {c.app_version}
                        """)


    ### MAIN PAGE

    # Title
    st.markdown("""# Ragnar
### Din GDPR- och sekretessäkrade transkriberare
""")
    st.markdown(f"""**Vald AI-modell:** {st.session_state["transcribe_model"]}   
**Valt språk:** {st.session_state["spoken_language"]}""")


    # CREATE TWO TABS FOR FILE UPLOAD VS RECORDED AUDIO    

    tab1, tab2 = st.tabs(["Ladda upp", "Spela in"])


    # FILE UPLOADER

    with tab1:
        
        uploaded_file = st.file_uploader(
            "Ladda upp din ljud- eller videofil här",
            type=["mp3", "wav", "flac", "mp4", "m4a", "aifc"],
            help="Max 2GB stora filer", label_visibility="collapsed",
            )


        if uploaded_file:

            # Checks if uploaded file has already been transcribed
            current_file_hash = compute_file_hash(uploaded_file)

            # If the uploaded file hash is different from the one in session state, reset the state
            if "file_hash" not in st.session_state or st.session_state.file_hash != current_file_hash:
                st.session_state.file_hash = current_file_hash
                
                if "transcribed" in st.session_state:
                    del st.session_state.transcribed

            
            # If audio has not been transcribed
            if "transcribed" not in st.session_state:

                # Sends audio to be converted to mp3 and compressed
                with st.spinner('Din ljudfil är lite stor. Jag ska bara komprimera den lite först...', show_time=True):
                    st.session_state.file_name_converted = convert_to_mono_and_compress(uploaded_file, uploaded_file.name)
                    st.success('Inspelning komprimerad och klar. Startar transkribering.')

               # Transcribes audio with Whisper
                with st.spinner('Transkriberar. Det här kan ta ett litet tag beroende på hur lång inspelningen är...', show_time=True):
                    
                    if "KB" in st.session_state["transcribe_model"]:
                        st.session_state.transcribed = transcribe_with_kb_whisper(st.session_state.file_name_converted, 
                            uploaded_file.name, 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            model_map_spoken_language[st.session_state["spoken_language"]])
                    else:
                        st.session_state.transcribed = transcribe_with_whisper(st.session_state.file_name_converted, 
                            uploaded_file.name, 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            model_map_spoken_language[st.session_state["spoken_language"]])
                    
                    st.success('Transkribering klar.')

                    st.balloons()
                    

            # Creates a Word document with the transcribed text
            document = Document()
            document.add_paragraph(st.session_state.transcribed)

            document.save('text/' + uploaded_file.name + '.docx')

            with open("text/" + uploaded_file.name + ".docx", "rb") as template_file:
                template_byte = template_file.read()


            # Creates a grid of four columns for the different transcribed document download buttons
            col1, col2, col3, col4 = st.columns(4)
            
            # Text
            with col1:
                with open('text/' + uploaded_file.name + '.txt', "rb") as file_txt:
                    st.download_button(
                        label = ":flag-se: Ladda ned text",
                        data = file_txt,
                        file_name = uploaded_file.name + '.txt',
                        mime = 'text/plain',
                    )

            # Word
            with col2:
                st.download_button(
                    label = ":flag-se: Ladda ned word",
                    data = template_byte,
                    file_name = uploaded_file.name + '.docx',
                    mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                )
            
            st.markdown("### Transkribering")
            
            if st.session_state.file_name_converted is not None:
                st.audio(st.session_state.file_name_converted, format='audio/wav')
            
            st.write(st.session_state.transcribed)

    
    # AUDIO RECORDER ###### ###### ######

    with tab2:

        audio = st.audio_input("Spela in")

        # The rest of the code in tab2 works the same way as in tab1, so it's not going to be
        # commented.
        if audio:

            # Open the saved audio file and compute its hash
            current_file_hash = compute_file_hash(audio)

            # If the uploaded file hash is different from the one in session state, reset the state
            if "file_hash" not in st.session_state or st.session_state.file_hash != current_file_hash:
                st.session_state.file_hash = current_file_hash
                
                if "transcribed" in st.session_state:
                    del st.session_state.transcribed

            if "transcribed" not in st.session_state:

                audio_file = AudioSegment.from_file(audio)
                output_path = "audio/converted.mp3"
                audio_file.export(output_path, format="mp3", bitrate="16k")
                    
                with st.spinner('Transkriberar. Det här kan ta ett litet tag beroende på hur lång inspelningen är...', show_time=True):
                    
                    if "KB" in st.session_state["transcribe_model"]:
                        st.session_state.transcribed = transcribe_with_kb_whisper("audio/converted.mp3", 
                            "local_recording.mp3", 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            model_map_spoken_language[st.session_state["spoken_language"]])
                    else:
                        st.session_state.transcribed = transcribe_with_whisper("audio/converted.mp3", 
                            "local_recording.mp3", 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            model_map_spoken_language[st.session_state["spoken_language"]])

                    st.success('Transkribering klar.')

                    st.balloons()

            local_recording_name = "local_recording.mp3"
            document = Document()
            document.add_paragraph(st.session_state.transcribed)

            document.save('text/' + local_recording_name + '.docx')

            with open("text/local_recording.mp3.docx", "rb") as template_file:
                template_byte = template_file.read()

            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                with open('text/' + local_recording_name + '.txt', "rb") as file_txt:
                    st.download_button(
                        label = ":flag-se: Ladda ned text",
                        data = file_txt,
                        file_name = local_recording_name + '.txt',
                        mime = 'text/plain',
                    )

            with col2:
                st.download_button(
                    label = ":flag-se: Ladda ned word",
                    data = template_byte,
                    file_name = local_recording_name + '.docx',
                    mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                )

            
            st.markdown("### Transkribering")
            
            if st.session_state.file_name_converted is not None:
                st.audio(st.session_state.file_name_converted, format='audio/wav')
            
            st.write(st.session_state.transcribed)


if __name__ == "__main__":
    main()